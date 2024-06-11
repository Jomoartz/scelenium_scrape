import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from docx import Document
from docx.shared import Pt

# Set up Chrome options
chrome_options = Options()
chrome_options.add_argument("--headless")  # Run in headless mode
chrome_options.add_argument("--disable-gpu")  # Disable GPU acceleration

# Set up the WebDriver
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)

# URL of the webpage to scrape
url = 'https://books.toscrape.com/'

# Open the webpage
driver.get(url)

# List to hold book data
books_data = []

# Function to scrape data from a page
def scrape_page():
    products = driver.find_elements(By.CLASS_NAME, 'product_pod')

    for product in products:
        # Get image URL
        image_url = product.find_element(By.TAG_NAME, 'img').get_attribute('src')

        # Get book title
        title = product.find_element(By.TAG_NAME, 'h3').text

        # Get price
        price = product.find_element(By.CLASS_NAME, 'price_color').text

        # Get stock status
        stock_status = product.find_element(By.CLASS_NAME, 'availability').text.strip()

        # Append the data to the list
        books_data.append({
            'image_url': image_url,
            'title': title,
            'price': price,
            'stock_status': stock_status
        })

# Loop through all pages
while True:
    scrape_page()
    # Find the 'next' button and move to the next page if it exists
    try:
        next_button = driver.find_element(By.CLASS_NAME, 'next')
        next_button_link = next_button.find_element(By.TAG_NAME, 'a').get_attribute('href')
        driver.get(next_button_link)
    except:
        break  # Break the loop if there's no next button

# Close the browser
driver.quit()

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('Books Data', level=1)

# Add book data to the document
for book in books_data:
    doc.add_heading(book['title'], level=2)
    doc.add_paragraph(f"Price: {book['price']}")
    doc.add_paragraph(f"Stock Status: {book['stock_status']}")
    doc.add_paragraph(f"Image URL: {book['image_url']}")
    doc.add_paragraph("\n")

# Save the document
doc.save('books_data.docx')

print("Data has been written to books_data.docx")
